import io
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import docx2txt
from bs4 import BeautifulSoup
import aspose.words as aw




pytesseract.pytesseract.tesseract_cmd="C:\Program Files\Tesseract-OCR\\tesseract.exe"
class DocumentConverter:
    def extract_text_from_pdf(self, pdf_path):
        pages = convert_from_path(pdf_path, 500)
        text_data = ''
        for page in pages:
            text = pytesseract.image_to_string(page)
            text_data += text + '\n'
        return text_data

    def extract_text_from_docx(self, docx_path):
        doc = Document(docx_path)
        text_data = ''
        for para in doc.paragraphs:
            text_data += para.text + '\n'
        return text_data

    def extract_text_from_html(self, html_path):
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        soup = BeautifulSoup(html_content, 'html.parser')

        def extract_text_recursive(element):
            text = ''
            if element.name in ['script', 'style', 'head', 'title', 'meta']:
                return ''
            if isinstance(element, str):
                return element.strip() + '\n'
            for child in element.children:
                text += extract_text_recursive(child)
            return text

        text_data = extract_text_recursive(soup)
        return text_data


    def extract_text_from_doc(self, doc_path):
        doc = aw.Document(doc_path)
    
        # Initialize an empty string to store the text
        text = ""
        
        # Iterate over all paragraphs in the document and extract text
        for para in doc.get_child_nodes(aw.NodeType.PARAGRAPH, True):
            text += para.get_text() + "\n"
        
        return text
    def abc(self,b):
        print("a",b)

# Example usage:
converter = DocumentConverter()
# text_from_pdf = converter.extract_text_from_pdf("D:\Mined 24\Resume\Resume(1).pdf")
# print(text_from_pdf)

# text_from_docx = converter.convert_doc_to_txt("D:\Mined 24\Resume\Srikanth_SrDataEngineer - Copy.doc")
# print(text_from_docx)

# text_from_html = converter.extract_text_from_html('practical2a.html')
# print(text_from_html)

# print(converter.extract_text_from_pdf("D:\Mined 24\Resume\Resume(1).pdf"))
# converter.abc(22)
